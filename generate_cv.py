import os
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable, Image as RLImage
from PIL import Image as PILImage, ImageDraw
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def prepare_circular_or_rounded_photo(src_path, dest_path, size=(300, 360), radius=20):
    os.makedirs(os.path.dirname(dest_path), exist_ok=True)
    img = PILImage.open(src_path).convert("RGBA")
    
    # Crop to aspect ratio 300:360 (5:6) centered around head
    w, h = img.size
    target_ratio = 300 / 360
    current_ratio = w / h
    
    if current_ratio > target_ratio:
        new_w = int(h * target_ratio)
        left = (w - new_w) // 2
        img_cropped = img.crop((left, 0, left + new_w, h))
    else:
        new_h = int(w / target_ratio)
        # Top-weighted crop for portrait headshots
        img_cropped = img.crop((0, 0, w, new_h))
        
    img_resized = img_cropped.resize(size, PILImage.Resampling.LANCZOS)
    
    # Rounded corners mask
    mask = PILImage.new('L', size, 0)
    draw = ImageDraw.Draw(mask)
    draw.rounded_rectangle([(0, 0), size], radius=radius, fill=255)
    
    img_resized.putalpha(mask)
    
    # Save as PNG with alpha
    img_resized.save(dest_path, "PNG")
    return dest_path

def generate_pdf(output_paths, photo_path=None):
    for output_path in output_paths:
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc = SimpleDocTemplate(
            output_path,
            pagesize=letter,
            leftMargin=0.45 * inch,
            rightMargin=0.45 * inch,
            topMargin=0.4 * inch,
            bottomMargin=0.4 * inch
        )

        styles = getSampleStyleSheet()
        
        # Color Palette
        PRIMARY = colors.HexColor('#0f172a')     # Dark Slate Navy
        ACCENT = colors.HexColor('#0284c7')      # Ocean Blue
        MUTED = colors.HexColor('#475569')       # Muted Gray
        DARK_TEXT = colors.HexColor('#1e293b')   # Charcoal Body
        LINE_COLOR = colors.HexColor('#cbd5e1')  # Light Gray Divider
        
        # Custom Typography Styles
        name_style = ParagraphStyle(
            'NameStyle',
            fontName='Helvetica-Bold',
            fontSize=18,
            leading=21,
            textColor=PRIMARY,
            alignment=0 # Left
        )
        
        title_style = ParagraphStyle(
            'TitleStyle',
            fontName='Helvetica-Bold',
            fontSize=10.5,
            leading=13,
            textColor=ACCENT,
            alignment=0
        )
        
        contact_style = ParagraphStyle(
            'ContactStyle',
            fontName='Helvetica',
            fontSize=8,
            leading=11.5,
            textColor=MUTED,
            alignment=0
        )
        
        section_heading = ParagraphStyle(
            'SectionHeading',
            fontName='Helvetica-Bold',
            fontSize=10,
            leading=12,
            textColor=PRIMARY,
            spaceBefore=6,
            spaceAfter=2,
            keepWithNext=True
        )
        
        body_style = ParagraphStyle(
            'Body',
            fontName='Helvetica',
            fontSize=8.2,
            leading=11,
            textColor=DARK_TEXT
        )
        
        bullet_style = ParagraphStyle(
            'BulletText',
            fontName='Helvetica',
            fontSize=8.2,
            leading=11,
            textColor=DARK_TEXT,
            leftIndent=8,
            firstLineIndent=-6
        )
        
        job_title_style = ParagraphStyle(
            'JobTitle',
            fontName='Helvetica-Bold',
            fontSize=9,
            leading=11.5,
            textColor=DARK_TEXT
        )
        
        sub_title_style = ParagraphStyle(
            'SubTitle',
            fontName='Helvetica-Oblique',
            fontSize=8,
            leading=10.5,
            textColor=MUTED
        )
        
        date_style = ParagraphStyle(
            'DateText',
            fontName='Helvetica-Bold',
            fontSize=8.2,
            leading=11,
            textColor=ACCENT,
            alignment=2 # Right
        )
        
        story = []
        
        # 1. HEADER (Left: Info, Right: Photo)
        header_text_elements = [
            Paragraph("CHANNA KAVISHKA SADARUWAN", name_style),
            Spacer(1, 2),
            Paragraph("Software Developer | Full-Stack &amp; 3D WebXR Engineer", title_style),
            Spacer(1, 3),
            Paragraph("Kandy, Sri Lanka  •  +94 70 457 3602  •  <a href='mailto:channasadhruvan@gmail.com' color='#0284c7'>channasadhruvan@gmail.com</a>", contact_style),
            Paragraph("<a href='https://github.com/C-KAVISHKA' color='#0284c7'>github.com/C-KAVISHKA</a>  •  <a href='https://linkedin.com/in/channa-sandaruwan' color='#0284c7'>linkedin.com/in/channa-sandaruwan</a>", contact_style)
        ]
        
        if photo_path and os.path.exists(photo_path):
            photo_img = RLImage(photo_path, width=0.95 * inch, height=1.14 * inch)
            header_table_data = [[header_text_elements, photo_img]]
            header_table = Table(header_table_data, colWidths=[6.4 * inch, 1.1 * inch])
            header_table.setStyle(TableStyle([
                ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
                ('ALIGN', (1,0), (1,0), 'RIGHT'),
                ('LEFTPADDING', (0,0), (-1,-1), 0),
                ('RIGHTPADDING', (0,0), (-1,-1), 0),
                ('TOPPADDING', (0,0), (-1,-1), 0),
                ('BOTTOMPADDING', (0,0), (-1,-1), 0),
            ]))
            story.append(header_table)
        else:
            for elem in header_text_elements:
                story.append(elem)
                
        story.append(HRFlowable(width="100%", thickness=1, color=LINE_COLOR, spaceBefore=4, spaceAfter=4))
        
        # 2. PROFESSIONAL PROFILE
        story.append(Paragraph("PROFESSIONAL SUMMARY", section_heading))
        story.append(HRFlowable(width="100%", thickness=0.75, color=ACCENT, spaceBefore=1, spaceAfter=3))
        profile_text = (
            "Final-year <b>BSc (Hons) Software Engineering</b> student at Cardiff Metropolitan University (graduating in 2026). "
            "Skilled in <b>Java, JavaScript (ES6+), TypeScript, and Python</b>, with hands-on experience building full-stack web platforms "
            "using the <b>MERN Stack</b> (MongoDB, Express.js, React, Node.js), <b>Java Spring Boot</b> with MySQL, and interactive <b>3D WebGL / WebXR</b> (Three.js). "
            "Seeking a trainee or software developer position to contribute to scalable software solutions and high-impact engineering teams."
        )
        story.append(Paragraph(profile_text, body_style))
        story.append(Spacer(1, 4))
        
        # 3. TECHNICAL SKILLS
        story.append(Paragraph("TECHNICAL SKILLS", section_heading))
        story.append(HRFlowable(width="100%", thickness=0.75, color=ACCENT, spaceBefore=1, spaceAfter=3))
        
        skills_data = [
            [Paragraph("<b>Programming Languages:</b>", body_style), Paragraph("Java (Core &amp; OOP), JavaScript (ES6+), TypeScript, Python, SQL, HTML5, CSS3", body_style)],
            [Paragraph("<b>Full-Stack &amp; Web:</b>", body_style), Paragraph("React.js, Next.js (App Router), Node.js, Express.js, Tailwind CSS, Vite, RESTful APIs", body_style)],
            [Paragraph("<b>3D &amp; Immersive Tech:</b>", body_style), Paragraph("Three.js, @react-three/fiber, @react-three/drei, WebXR (AR/VR room placement)", body_style)],
            [Paragraph("<b>Backend &amp; Databases:</b>", body_style), Paragraph("Spring Boot, MySQL, MongoDB, Mongoose, Hibernate/JPA, Stripe API, Cloudinary", body_style)],
            [Paragraph("<b>Tools &amp; Workflow:</b>", body_style), Paragraph("Git, GitHub, Postman, Maven, VS Code, Linux/Windows CLI, Agile/Scrum", body_style)]
        ]
        skills_table = Table(skills_data, colWidths=[1.6 * inch, 5.9 * inch])
        skills_table.setStyle(TableStyle([
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('BOTTOMPADDING', (0,0), (-1,-1), 1),
            ('TOPPADDING', (0,0), (-1,-1), 1),
            ('LEFTPADDING', (0,0), (-1,-1), 0),
            ('RIGHTPADDING', (0,0), (-1,-1), 0),
        ]))
        story.append(skills_table)
        story.append(Spacer(1, 4))
        
        # 4. KEY PROJECTS
        story.append(Paragraph("FEATURED SOFTWARE PROJECTS", section_heading))
        story.append(HRFlowable(width="100%", thickness=0.75, color=ACCENT, spaceBefore=1, spaceAfter=3))
        
        # Project 1: iFurnish Shop
        p1_header = [
            [Paragraph("<b>iFurnish Shop</b> — 3D WebXR Furniture E-Commerce Platform", job_title_style), Paragraph("<b>MERN Stack | Three.js | WebXR</b>", date_style)]
        ]
        t1 = Table(p1_header, colWidths=[5.3 * inch, 2.2 * inch])
        t1.setStyle(TableStyle([('VALIGN', (0,0), (-1,-1), 'MIDDLE'), ('LEFTPADDING', (0,0), (-1,-1), 0), ('RIGHTPADDING', (0,0), (-1,-1), 0), ('BOTTOMPADDING', (0,0), (-1,-1), 1)]))
        story.append(t1)
        story.append(Paragraph("• Architected a production-grade 3D e-commerce web platform enabling real-time WebGL model customization (colors, materials) and browser-based Augmented Reality (WebXR) room projection.", bullet_style))
        story.append(Paragraph("• Engineered complete backend services with Node.js, Express, and MongoDB, integrating secure Stripe checkout payments, Cloudinary CDN media uploads, and JWT authentication.", bullet_style))
        story.append(Paragraph("• Developed a comprehensive Admin Dashboard for dynamic product inventory management and real-time sales tracking.", bullet_style))
        story.append(Spacer(1, 3))
        
        # Project 2: Anime Site / Onisaga
        p2_header = [
            [Paragraph("<b>Anime Site / Onisaga Media Hub</b> — Dynamic Content Discovery Web App", job_title_style), Paragraph("<b>MERN Stack | Next.js</b>", date_style)]
        ]
        t2 = Table(p2_header, colWidths=[5.3 * inch, 2.2 * inch])
        t2.setStyle(TableStyle([('VALIGN', (0,0), (-1,-1), 'MIDDLE'), ('LEFTPADDING', (0,0), (-1,-1), 0), ('RIGHTPADDING', (0,0), (-1,-1), 0), ('BOTTOMPADDING', (0,0), (-1,-1), 1)]))
        story.append(t2)
        story.append(Paragraph("• Built a responsive media discovery web application with dynamic catalog browsing, real-time search queries, and interactive carousel sliders.", bullet_style))
        story.append(Paragraph("• Implemented RESTful API endpoints for streaming metadata and user bookmark lists with fast MongoDB document lookups.", bullet_style))
        story.append(Spacer(1, 3))
        
        # Project 3: Oceanview Reservation System
        p3_header = [
            [Paragraph("<b>Oceanview Reservation System</b> — Enterprise Hotel &amp; Resort Booking Platform", job_title_style), Paragraph("<b>Java Spring Boot | MySQL</b>", date_style)]
        ]
        t3 = Table(p3_header, colWidths=[5.3 * inch, 2.2 * inch])
        t3.setStyle(TableStyle([('VALIGN', (0,0), (-1,-1), 'MIDDLE'), ('LEFTPADDING', (0,0), (-1,-1), 0), ('RIGHTPADDING', (0,0), (-1,-1), 0), ('BOTTOMPADDING', (0,0), (-1,-1), 1)]))
        story.append(t3)
        story.append(Paragraph("• Developed a multi-tier reservation management system enabling automated room bookings, availability management, and billing calculations.", bullet_style))
        story.append(Paragraph("• Designed normalized MySQL database schemas with ACID transaction compliance, foreign key constraints, and role-based access control.", bullet_style))
        story.append(Spacer(1, 3))
        
        # Project 4: HealthShield AI
        p4_header = [
            [Paragraph("<b>HealthShield AI</b> — Health Record &amp; Diagnostics Management Application", job_title_style), Paragraph("<b>Java Spring Boot | MySQL</b>", date_style)]
        ]
        t4 = Table(p4_header, colWidths=[5.3 * inch, 2.2 * inch])
        t4.setStyle(TableStyle([('VALIGN', (0,0), (-1,-1), 'MIDDLE'), ('LEFTPADDING', (0,0), (-1,-1), 0), ('RIGHTPADDING', (0,0), (-1,-1), 0), ('BOTTOMPADDING', (0,0), (-1,-1), 1)]))
        story.append(t4)
        story.append(Paragraph("• Developed backend RESTful services for secure patient health metric logging, appointment scheduling, and record analytics.", bullet_style))
        story.append(Paragraph("• Integrated secure data persistence layer in MySQL with data validation and role-based authorization.", bullet_style))
        story.append(Spacer(1, 4))
        
        # 5. EDUCATION
        story.append(Paragraph("EDUCATION &amp; QUALIFICATIONS", section_heading))
        story.append(HRFlowable(width="100%", thickness=0.75, color=ACCENT, spaceBefore=1, spaceAfter=3))
        
        edu_data = [
            [
                Paragraph("<b>BSc (Hons) in Software Engineering</b>", job_title_style),
                Paragraph("<b>Expected 2026</b>", date_style)
            ],
            [
                Paragraph("Cardiff Metropolitan University (Kandy Campus, Sri Lanka)", sub_title_style),
                Paragraph("", sub_title_style)
            ],
            [
                Paragraph("<b>Higher National Diploma (HND) in Software Engineering</b>", job_title_style),
                Paragraph("<b>2021 — 2023</b>", date_style)
            ],
            [
                Paragraph("ICBT Campus, Sri Lanka", sub_title_style),
                Paragraph("", sub_title_style)
            ],
            [
                Paragraph("<b>Diploma in English</b> — British Way English Academy", body_style),
                Paragraph("<b>Completed</b>", date_style)
            ],
            [
                Paragraph("<b>G.C.E. Advanced Level (A/L) &amp; Ordinary Level (O/L)</b>", body_style),
                Paragraph("<b>Completed</b>", date_style)
            ]
        ]
        edu_table = Table(edu_data, colWidths=[5.7 * inch, 1.8 * inch])
        edu_table.setStyle(TableStyle([
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('LEFTPADDING', (0,0), (-1,-1), 0),
            ('RIGHTPADDING', (0,0), (-1,-1), 0),
            ('TOPPADDING', (0,0), (-1,-1), 1),
            ('BOTTOMPADDING', (0,0), (-1,-1), 1),
        ]))
        story.append(edu_table)
        story.append(Spacer(1, 4))
        
        # 6. REFERENCES
        story.append(Paragraph("REFERENCES", section_heading))
        story.append(HRFlowable(width="100%", thickness=0.75, color=ACCENT, spaceBefore=1, spaceAfter=2))
        story.append(Paragraph("Available upon request.", body_style))
        
        doc.build(story)
        print(f"Generated PDF with Photo at: {output_path}")

def generate_docx(output_path, photo_path=None):
    doc = docx.Document()
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
        
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_name = p_name.add_run("CHANNA KAVISHKA SADARUWAN")
    run_name.font.name = 'Calibri'
    run_name.font.size = Pt(20)
    run_name.font.bold = True
    run_name.font.color.rgb = RGBColor(15, 23, 42)
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("Software Developer | Full-Stack & 3D WebXR Engineer")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(11)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(2, 132, 199)
    
    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_contact.paragraph_format.space_after = Pt(8)
    run_contact = p_contact.add_run("Kandy, Sri Lanka  |  +94 70 457 3602  |  channasadhruvan@gmail.com\ngithub.com/C-KAVISHKA  |  linkedin.com/in/channa-sandaruwan")
    run_contact.font.name = 'Calibri'
    run_contact.font.size = Pt(9.5)
    run_contact.font.color.rgb = RGBColor(71, 85, 105)
    
    def add_section_header(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(title)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(15, 23, 42)
        
    add_section_header("PROFESSIONAL SUMMARY")
    p_prof = doc.add_paragraph()
    p_prof.paragraph_format.space_after = Pt(6)
    r_prof = p_prof.add_run("Final-year BSc (Hons) Software Engineering student at Cardiff Metropolitan University (Expected 2026). Skilled in Java, JavaScript (ES6+), TypeScript, and Python, with hands-on experience building full-stack web platforms using the MERN Stack (MongoDB, Express.js, React, Node.js), Java Spring Boot with MySQL, and interactive 3D WebGL / WebXR (Three.js). Seeking a trainee or software developer position in a high-impact engineering environment.")
    r_prof.font.name = 'Calibri'
    r_prof.font.size = Pt(9.5)
    
    add_section_header("TECHNICAL SKILLS")
    skills_list = [
        ("Programming Languages: ", "Java (Core & OOP), JavaScript (ES6+), TypeScript, Python, SQL, HTML5, CSS3"),
        ("Full-Stack & Web: ", "React.js, Next.js (App Router), Node.js, Express.js, Tailwind CSS, Vite, RESTful APIs"),
        ("3D & Immersive Tech: ", "Three.js, @react-three/fiber, @react-three/drei, WebXR (AR/VR)"),
        ("Backend & Databases: ", "Spring Boot, MySQL, MongoDB, Mongoose, Hibernate/JPA, Stripe API, Cloudinary"),
        ("Tools & Workflow: ", "Git, GitHub, Postman, Maven, VS Code, Agile/Scrum")
    ]
    for cat, items in skills_list:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(1.5)
        r1 = p.add_run(cat)
        r1.font.bold = True
        r1.font.size = Pt(9.5)
        r2 = p.add_run(items)
        r2.font.size = Pt(9.5)
        
    add_section_header("FEATURED SOFTWARE PROJECTS")
    
    projects = [
        ("iFurnish Shop — 3D WebXR Furniture E-Commerce Platform", "MERN Stack, Three.js, WebXR", [
            "Architected a production-grade 3D e-commerce platform enabling real-time WebGL model customization and mobile browser Augmented Reality (WebXR) room projection.",
            "Engineered complete backend services with Node.js, Express, MongoDB, Stripe payment checkout, and Cloudinary media pipelines.",
            "Developed an Admin Dashboard for real-time inventory management and sales analytics."
        ]),
        ("Anime Site / Onisaga Media Hub — Dynamic Content Discovery Web App", "MERN Stack, Next.js", [
            "Built a responsive media discovery web application with dynamic catalog browsing, real-time search queries, and interactive carousel sliders.",
            "Implemented RESTful API endpoints for streaming metadata and user bookmark lists with fast MongoDB document lookups."
        ]),
        ("Oceanview Reservation System — Enterprise Hotel Booking Platform", "Java Spring Boot, MySQL", [
            "Developed a multi-tier reservation system enabling automated room bookings, availability management, and billing calculations.",
            "Designed normalized MySQL relational schemas with ACID transaction compliance, foreign key constraints, and role-based access control."
        ]),
        ("HealthShield AI — Health Record & Diagnostics Management Application", "Java Spring Boot, MySQL", [
            "Developed backend RESTful services for secure patient health metric logging, appointment scheduling, and record analytics.",
            "Integrated secure data persistence layer in MySQL with data validation and role-based authorization."
        ])
    ]
    
    for title, tech, bullets in projects:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(1)
        r_t = p.add_run(title)
        r_t.font.bold = True
        r_t.font.size = Pt(9.5)
        r_sub = p.add_run(f" | {tech}")
        r_sub.font.italic = True
        r_sub.font.size = Pt(9)
        r_sub.font.color.rgb = RGBColor(2, 132, 199)
        
        for b in bullets:
            pb = doc.add_paragraph(style='List Bullet')
            pb.paragraph_format.space_after = Pt(1)
            rb = pb.add_run(b)
            rb.font.size = Pt(9.5)
            
    add_section_header("EDUCATION & QUALIFICATIONS")
    
    edus = [
        ("BSc (Hons) in Software Engineering", "Expected 2026", "Cardiff Metropolitan University (Kandy Campus, Sri Lanka)"),
        ("Higher National Diploma (HND) in Software Engineering", "2021 — 2023", "ICBT Campus, Sri Lanka"),
        ("Diploma in English", "Completed", "British Way English Academy"),
        ("G.C.E. Advanced Level (A/L) & Ordinary Level (O/L)", "Completed", "Secondary Education")
    ]
    for deg, yr, inst in edus:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(1)
        rd = p.add_run(deg)
        rd.font.bold = True
        rd.font.size = Pt(9.5)
        ry = p.add_run(f" ({yr})")
        ry.font.bold = True
        ry.font.color.rgb = RGBColor(2, 132, 199)
        ry.font.size = Pt(9)
        
        pi = doc.add_paragraph()
        pi.paragraph_format.space_after = Pt(3)
        ri = pi.add_run(inst)
        ri.font.italic = True
        ri.font.size = Pt(9)
        ri.font.color.rgb = RGBColor(71, 85, 105)
        
    add_section_header("REFERENCES")
    p_ref = doc.add_paragraph()
    r_ref = p_ref.add_run("Available upon request.")
    r_ref.font.size = Pt(9.5)
    
    doc.save(output_path)
    print(f"Generated DOCX at: {output_path}")

if __name__ == "__main__":
    raw_photo = r"C:\Users\Enzo\.gemini\antigravity-ide\scratch\portfolio\public\profile.jpg"
    processed_photo = r"C:\Users\Enzo\.gemini\antigravity-ide\scratch\portfolio\public\profile_headshot.png"
    prepare_circular_or_rounded_photo(raw_photo, processed_photo)
    
    pdf_paths = [
        r"C:\Users\Enzo\.gemini\antigravity-ide\scratch\portfolio\public\Channa_Kavishka_CV.pdf",
        r"C:\Users\Enzo\Downloads\Channa_Kavishka_CV.pdf"
    ]
    generate_pdf(pdf_paths, photo_path=processed_photo)
    
    docx_path = r"C:\Users\Enzo\Downloads\Channa_Kavishka_CV_Professional.docx"
    generate_docx(docx_path)
